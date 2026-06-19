"""Generate graduation final report as Word (.docx) — English version."""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm, Inches

ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / 'docs' / 'Final_Report_Jadah_Hail_Final.docx'
DIAGRAMS = ROOT / 'docs' / 'diagrams'
SCREENSHOTS = ROOT / 'docs' / 'screenshots'


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(16 if level == 1 else 14)
    return h


def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(item, style='List Bullet')
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.name = 'Times New Roman'
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
            for p in cells[i].paragraphs:
                for run in p.runs:
                    run.font.name = 'Times New Roman'
    doc.add_paragraph()


def add_image(doc, path, caption=None, width=Inches(6.2)):
    if not path.exists():
        add_para(doc, f'[Diagram not found: {path.name}]')
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=width)
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run(caption)
        r.italic = True
        r.font.name = 'Times New Roman'
        r.font.size = Pt(10)
    doc.add_paragraph()


def build_document():
    doc = Document()
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # Title page
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run('Graduation Project\n\n')
    r.bold = True
    r.font.size = Pt(20)
    r.font.name = 'Times New Roman'
    r2 = t.add_run('Jadah Hail — Smart Tourism Platform\n\n')
    r2.font.size = Pt(18)
    r2.font.name = 'Times New Roman'
    r3 = t.add_run(
        'Prepared by: Rahaf Bader Alzabni\n'
        'Student ID: s202005878\n\n'
        'Hail Municipality\n'
        '2025'
    )
    r3.font.size = Pt(14)
    r3.font.name = 'Times New Roman'
    doc.add_page_break()

    # Abstract
    add_heading(doc, 'Abstract')
    add_para(doc,
        'This graduation project presents Jadah Hail, a smart tourism platform developed for '
        'Hail Municipality in the Kingdom of Saudi Arabia. The system integrates a React-based '
        'visitor portal with a Django REST API backend to showcase tourist landmarks, events, '
        'suggested routes, interactive maps, user reviews, favorites, and a rule-based tourism '
        'assistant. The platform supports bilingual Arabic/English interfaces with RTL layout. '
        'Testing includes 19 automated Django unit tests covering places, reviews, routes, '
        'events, authentication, and favorites APIs. Results demonstrate a functional, '
        'scalable client-server architecture suitable for smart-city tourism initiatives.')
    doc.add_page_break()

    # 1. Users and Project Scope
    add_heading(doc, '1. Users and Project Scope')
    add_para(doc,
        'The Jadah Hail project aims to build a smart tourism platform for the Hail region. '
        'It combines tourist attractions, events, suggested routes, interactive maps, reviews, '
        'and a tourism assistant, along with an admin dashboard for municipality staff.')
    add_para(doc, 'Project scope includes:', bold=True)
    add_bullets(doc, [
        'A visitor portal supporting Arabic and English languages.',
        'Display and exploration of tourist sites in Hail with images, descriptions, and coordinates.',
        'A real interactive map (OpenStreetMap + Leaflet) showing landmark locations.',
        'Display of tourism events and suggested travel routes.',
        'A review and favorites system for registered visitors.',
        'A rule-based tourism assistant.',
        'An admin dashboard for viewing statistics and managing content (UI layer).',
        'A REST API built with Django REST Framework.',
    ])
    add_para(doc, 'Out of Scope:', bold=True)
    add_bullets(doc, [
        'Native mobile applications (iOS/Android) — can be added later using the same API.',
        'Electronic payments or hotel booking.',
        'Generative AI (e.g., ChatGPT) — the current assistant uses stored rule-based replies.',
    ])

    # 2. People who use the system
    add_heading(doc, '2. People Who Use the System')
    add_table(doc,
        ['User Type', 'Description', 'Permissions'],
        [
            ['Guest Visitor', 'Browses the platform without registration',
             'View places, events, routes, map, and assistant'],
            ['Registered Visitor', 'Has a Django user account',
             'All guest permissions + favorites + submit reviews'],
            ['Municipality Admin', 'Hail Municipality staff member',
             'Admin dashboard + Django Admin for data management'],
            ['Developer / IT Team', 'Maintains the system',
             'Server, database, and deployment management'],
        ])

    # 3. Requirement Analysis
    add_heading(doc, '3. Requirement Analysis')
    add_para(doc, '3.1 Functional Requirements', bold=True)
    add_bullets(doc, [
        'FR-01: Display a list of tourist sites with category (heritage, natural, cultural, etc.).',
        'FR-02: Search and filter places by category.',
        'FR-03: Detail page for each landmark (description, hours, map, reviews).',
        'FR-04: Interactive map of the Hail region with landmark markers.',
        'FR-05: Display upcoming tourism events.',
        'FR-06: Display suggested tourist routes.',
        'FR-07: Visitor login and logout.',
        'FR-08: Add/remove places from favorites.',
        'FR-09: Submit a rating (1–5) and comment for each place.',
        'FR-10: Tourism assistant answering common questions.',
        'FR-11: Admin dashboard for statistics, places, and events.',
        'FR-12: Support for Arabic (RTL) and English languages.',
    ])
    add_para(doc, '3.2 Non-Functional Requirements', bold=True)
    add_bullets(doc, [
        'NFR-01: User-friendly and responsive interface.',
        'NFR-02: API response time under 2 seconds in the local environment.',
        'NFR-03: Session security and CSRF protection.',
        'NFR-04: Separation of frontend and backend (Client-Server Architecture).',
        'NFR-05: Scalability for future deployment (PostgreSQL, cloud hosting).',
    ])

    # 4. Software Requirements
    add_heading(doc, '4. Software Requirements')
    add_table(doc,
        ['Software', 'Version', 'Purpose'],
        [
            ['Python', '3.13+', 'Backend language'],
            ['Django', '6.x', 'Web framework'],
            ['Django REST Framework', '3.15+', 'REST API'],
            ['django-cors-headers', '4.x', 'Allow frontend API access'],
            ['Pillow', '10+', 'Image processing for places and events'],
            ['Node.js', '24+', 'Frontend tooling'],
            ['React', '18.3', 'User interface'],
            ['TypeScript', '5.x', 'Type-safe frontend code'],
            ['Vite', '6.3', 'Frontend build and dev server'],
            ['Tailwind CSS', '4.1', 'UI styling'],
            ['Leaflet', '1.9', 'Interactive map'],
            ['Recharts', '2.15', 'Admin dashboard charts'],
            ['SQLite', '3', 'Database (development)'],
            ['Git', '2.x', 'Version control'],
        ])

    # 5. Backend Framework
    add_heading(doc, '5. Backend Framework')
    add_para(doc,
        'The server is built with Django 6 and Django REST Framework (DRF). '
        'The architecture is divided into the following apps:')
    add_bullets(doc, [
        'tourism: Tourist places, reviews, and routes.',
        'events: Tourism events.',
        'accounts: User profile, favorites, and authentication.',
        'ai_assistant: Tourism assistant (stored replies).',
        'dashboard: Reserved for future expansion.',
    ])
    add_para(doc, 'Main API endpoints (/api/v1/):')
    add_table(doc,
        ['Endpoint', 'Method', 'Description'],
        [
            ['/places/', 'GET', 'List tourist places'],
            ['/reviews/', 'GET/POST', 'View/add reviews'],
            ['/routes/', 'GET', 'Tourist routes'],
            ['/events/', 'GET', 'Events'],
            ['/favorites/', 'GET/POST/DELETE', 'Favorites'],
            ['/auth/login/', 'POST', 'Login'],
            ['/auth/logout/', 'POST', 'Logout'],
            ['/auth/me/', 'GET', 'Current user status'],
            ['/assistant/', 'GET', 'Assistant settings'],
            ['/assistant/ask/', 'POST', 'Ask the assistant'],
        ])

    # 6. Frontend Technologies
    add_heading(doc, '6. Frontend Technologies')
    add_para(doc,
        'The frontend is built with React, TypeScript, Vite, and Tailwind CSS. '
        'The design is inspired by Figma and includes:')
    add_bullets(doc, [
        'React Context (TourismContext) for data state management.',
        'API layer (src/api/) for communication with Django.',
        'HailMap component for the interactive map (Leaflet + OpenStreetMap).',
        'AIAssistantSection for the tourism assistant.',
        'RTL support and full localization (src/lib/locale.ts).',
        'Vite proxy for /api and /media to Django during development.',
    ])

    # 7. Database
    add_heading(doc, '7. Database')
    add_para(doc, 'SQLite3 is used in the development environment. Main tables:')
    add_table(doc,
        ['Table / Model', 'Key Fields', 'Relationships'],
        [
            ['TouristPlace', 'name_ar, name_en, description, category, lat/lng, image',
             '← Review, M2M Route, M2M Favorites'],
            ['Review', 'user, place, rating, comment', 'FK → User, TouristPlace'],
            ['TouristRoute', 'name, duration, difficulty', 'M2M → TouristPlace'],
            ['Event', 'title_ar/en, dates, location, image', '—'],
            ['UserProfile', 'language, interests, favorites', 'OneToOne → User'],
            ['AssistantReply', 'prompt/response ar/en, keywords', '—'],
        ])
    add_para(doc,
        'Figure 3 presents the Entity Relationship Diagram (ERD) showing primary keys, '
        'foreign keys, and many-to-many relationships between core entities.')
    add_image(doc, DIAGRAMS / 'erd_diagram.png',
              caption='Figure 3: Entity Relationship Diagram (ERD)')

    # 8. Development Environment
    add_heading(doc, '8. Development Environment')
    add_bullets(doc, [
        'Operating System: Windows 10/11',
        'IDE: Cursor / Visual Studio Code',
        'Backend: python manage.py runserver → http://127.0.0.1:8000',
        'Frontend: npm run dev → http://localhost:5173',
        'Django Admin: http://127.0.0.1:8000/admin/',
        'Data commands: load_demo_data, sync_arabic_content, import_hail_images, load_assistant_data',
    ])

    # 9. Version Control
    add_heading(doc, '9. Version Control')
    add_para(doc, 'Git + GitHub. Repository: https://github.com/moneerafahaid-collab/WATN-2-.git')
    add_bullets(doc, [
        'Commit messages describe the feature or fix.',
        'Branches used for development when needed.',
        'Backend and frontend changes tracked together.',
    ])

    # 10. Hardware Requirements
    add_heading(doc, '10. Hardware Requirements')
    add_table(doc,
        ['Component', 'Minimum', 'Recommended'],
        [
            ['Processor', 'Intel i3 / Ryzen 3', 'Intel i5 / Ryzen 5'],
            ['RAM', '8 GB', '16 GB'],
            ['Storage', '10 GB free', '20 GB SSD'],
            ['Display', '1366×768', '1920×1080'],
            ['Network', 'Internet (for map tiles and images)', 'Broadband'],
        ])

    # 11. System Design and Flow
    add_heading(doc, '11. System Design and Flow')
    add_para(doc,
        'The platform follows a three-tier client-server architecture separating presentation, '
        'application logic, and data persistence layers.')
    add_image(doc, DIAGRAMS / 'architecture_diagram.png',
              caption='Figure 4: Three-Tier System Architecture')
    add_para(doc,
        'Data flow: User opens browser → React requests /api/v1/places/ → '
        'Django reads from SQLite → returns JSON → React renders the data. '
        'Map tiles are fetched from OpenStreetMap at runtime.')

    # 12. Use Case Diagram
    add_heading(doc, '12. Use Case Diagram')
    add_para(doc,
        'Figure 1 shows the main actors and use cases of the Jadah Hail platform. '
        'Three actors interact with the system: Visitor, Registered User, and Admin.')
    add_image(doc, DIAGRAMS / 'use_case_diagram.png',
              caption='Figure 1: Use Case Diagram')
    add_para(doc, 'Key relationships:', bold=True)
    add_bullets(doc, [
        'Visitor: Browse Places, View Map, View Events, View Routes, Use Assistant, Switch Language.',
        'Registered User: Login, Add Favorite, Submit Review (+ all Visitor permissions).',
        'Admin: View Dashboard, Manage Content, View Analytics.',
        'Submit Review <<include>> Login.',
        'View Place Details <<extend>> Browse Places.',
    ])

    # 13. Sequence Diagram
    add_heading(doc, '13. Sequence Diagram')
    add_para(doc,
        'Figure 2 illustrates the sequence of interactions when a registered user submits a review.')
    add_image(doc, DIAGRAMS / 'sequence_diagram.png',
              caption='Figure 2: Sequence Diagram — Submit Review')
    add_bullets(doc, [
        'User fills the review form and clicks Submit.',
        'React sends POST /api/v1/reviews/ with CSRF token and session cookie.',
        'Django validates IsAuthenticated permission.',
        'Django inserts a new Review record into SQLite.',
        'API returns JSON (201 Created) and React updates the UI.',
    ])

    # 14. Website Implementation
    add_heading(doc, '14. Website Implementation')
    add_para(doc, '14.1 Backend', bold=True)
    add_bullets(doc, [
        'Models in tourism, events, accounts, and ai_assistant apps.',
        'Serializers to convert data to JSON.',
        'ViewSets with AllowAny / IsAuthenticated permissions.',
        'Management commands to load Hail demo data and images.',
    ])
    add_para(doc, '14.2 Frontend', bold=True)
    add_bullets(doc, [
        'App.tsx: Home, Explore, Details, Map, Routes, Assistant, and Admin Dashboard.',
        'HailMap.tsx: Leaflet interactive map.',
        'TourismContext: Fetches data from the API.',
        'locale.ts + mappers.ts: Localization and data mapping.',
    ])

    # 15. Testing and Output Results
    add_heading(doc, '15. Testing and Output Results')
    add_para(doc, '15.1 Testing Types', bold=True)
    add_table(doc,
        ['Test Type', 'Method', 'Result'],
        [
            ['Manual functional testing', 'Browse all platform pages', 'Passed'],
            ['API integration testing', 'Postman / Browser DevTools', 'Passed'],
            ['Authentication testing', 'Session login/logout flow', 'Passed'],
            ['Map testing', 'Leaflet markers on OpenStreetMap', 'Passed'],
            ['Localization testing', 'Arabic RTL / English LTR switch', 'Passed'],
            ['Automated unit tests', 'python manage.py test (19 tests)', 'Passed — 19/19'],
        ])
    add_para(doc, '15.2 Automated Test Coverage', bold=True)
    add_table(doc,
        ['Test Module', 'Tests', 'Coverage Area'],
        [
            ['tourism.tests', '10', 'Places list/filter/search, reviews CRUD, routes'],
            ['accounts.tests', '7', 'Login, logout, session, favorites API'],
            ['events.tests', '2', 'Events list and bilingual fields'],
        ])
    add_para(doc,
        'Command: python manage.py test tourism accounts events — Result: OK (19 tests).')
    add_para(doc, '15.3 Expected Output', bold=True)
    add_para(doc,
        'The platform displays 6 tourist landmarks, 3 events, 3 routes, '
        'an assistant with stored replies, and a bilingual admin dashboard.')

    # 16. Web System Interface
    add_heading(doc, '16. Web System Interface')
    add_para(doc,
        'Figures 5–11 show screenshots of the deployed platform captured from the live '
        'development environment (React + Django).')
    screenshots = [
        ('01_home.png', 'Figure 5: Home Page'),
        ('02_explore.png', 'Figure 6: Explore Page'),
        ('03_map.png', 'Figure 7: Interactive Map'),
        ('04_routes.png', 'Figure 8: Tourist Routes'),
        ('05_assistant.png', 'Figure 9: AI Tourism Assistant'),
        ('06_place_details.png', 'Figure 10: Place Details'),
        ('07_admin_dashboard.png', 'Figure 11: Admin Dashboard'),
    ]
    for filename, caption in screenshots:
        add_image(doc, SCREENSHOTS / filename, caption=caption, width=Inches(6.0))

    # 17. Conclusion
    add_heading(doc, '17. Conclusion')
    add_para(doc,
        'The Jadah Hail platform was successfully implemented as a graduation project '
        'connecting the tourism heritage of the Hail region with modern technology. '
        'The project achieved its goals of displaying landmarks, events, and routes through '
        'a professional bilingual interface, with a structured API, interactive map, and review system.')
    add_para(doc, 'Future Recommendations:', bold=True)
    add_bullets(doc, [
        'Fully connect the admin dashboard to the API (CRUD operations).',
        'Expand automated test coverage to frontend (React Testing Library).',
        'Migrate to PostgreSQL and deploy on a production server.',
        'Develop a mobile application consuming the same API.',
        'Optionally integrate a real AI assistant (LLM).',
    ])

    # 18. Citations
    add_heading(doc, '18. Citations')
    add_bullets(doc, [
        'Django Software Foundation. (2025). Django Documentation. https://docs.djangoproject.com/',
        'Django REST Framework. (2025). API Guide. https://www.django-rest-framework.org/',
        'Meta Open Source. (2025). React Documentation. https://react.dev/',
        'Vite Team. (2025). Vite Guide. https://vite.dev/',
        'Leaflet. (2025). An open-source JavaScript library for mobile-friendly interactive maps. https://leafletjs.com/',
        'OpenStreetMap contributors. (2025). OpenStreetMap. https://www.openstreetmap.org/',
        'Wikimedia Commons. (2025). Images of Hail region landmarks. https://commons.wikimedia.org/',
        'UNESCO. Rock Art in the Hail Region of Saudi Arabia. https://whc.unesco.org/',
        'Tailwind Labs. (2025). Tailwind CSS Documentation. https://tailwindcss.com/docs',
        'Recharts. (2025). Composable charting library for React. https://recharts.org/',
    ])

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == '__main__':
    import subprocess
    import sys

    subprocess.run([sys.executable, str(ROOT / 'scripts' / 'generate_diagrams.py')], check=True)

    if not any(SCREENSHOTS.glob('*.png')):
        print('Note: No screenshots found. Run scripts/capture_screenshots.py with servers running.')

    path = build_document()
    print(f'Created: {path}')
