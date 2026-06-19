"""
Generate University of Hail COOP Final Report (.docx)
College of Computer Science and Engineering — COOP Education Office
"""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Cm, Inches

ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / 'docs' / 'Final_COOP_Report_Jadah_Hail_v2.docx'
DIAGRAMS = ROOT / 'docs' / 'diagrams'
SCREENSHOTS = ROOT / 'docs' / 'screenshots'

# ── Report metadata (edit faculty advisor name before submission) ──────────────
STUDENT_NAME = 'Rahaf Bader Alzabni'
STUDENT_ID = 's202005878'
DEPARTMENT = 'College of Computer Science and Engineering'
FACULTY_ADVISOR = '[Faculty Advisor Name]'
REPORT_DATE = 'June 2026'
EMPLOYER = 'Hail Municipality (Hail Region, Kingdom of Saudi Arabia)'
WORK_UNIT = (
    'Technology Department — Digital Transformation Agency '
    '(وكالة التحول الرقمي — قسم التقنية)'
)
REPORT_TITLE = (
    'Development of Jadah Hail Digital: A Smart Tourism Guide '
    'Integrated with Tawakkalna Municipality Services'
)

FONT = 'Times New Roman'
FONT_SIZE = Pt(12)


# ── Formatting helpers ───────────────────────────────────────────────────────

def style_run(run, bold=False, size=FONT_SIZE, italic=False):
    run.font.name = FONT
    run.font.size = size
    run.bold = bold
    run.italic = italic


def style_paragraph(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=6, line_spacing=1.5):
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = line_spacing


def add_para(doc, text, bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    style_run(run, bold=bold, italic=italic)
    style_paragraph(p, align=align)
    return p


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        style_run(run, bold=True, size=Pt(16 if level == 1 else 14))
    style_paragraph(h)
    return h


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(item, style='List Bullet')
        for run in p.runs:
            style_run(run)
        style_paragraph(p)


def add_table(doc, headers, rows, caption=None):
    if caption:
        add_para(doc, caption, bold=True)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                style_run(run, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
            for p in cells[i].paragraphs:
                for run in p.runs:
                    style_run(run)
    doc.add_paragraph()
    return table


def add_image(doc, path, caption=None, width=Inches(6.0)):
    if not path.exists():
        add_para(doc, f'[Figure not found: {path.name}]', italic=True)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER)
    p.add_run().add_picture(str(path), width=width)
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style_paragraph(cap, align=WD_ALIGN_PARAGRAPH.CENTER)
        style_run(cap.add_run(caption), italic=True, size=Pt(10))
    doc.add_paragraph()


def add_page_break(doc):
    doc.add_page_break()


def add_toc_line(doc, title, level=1, page=''):
    """One TOC entry with dot leader (page number optional)."""
    p = doc.add_paragraph()
    style_paragraph(p, line_spacing=1.5)
    if level > 1:
        p.paragraph_format.left_indent = Inches(0.35 * (level - 1))
    tab_pos = Inches(6.3)
    p.paragraph_format.tab_stops.add_tab_stop(
        tab_pos, WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
    run = p.add_run(title)
    style_run(run, bold=(level == 1))
    p.add_run('\t')
    style_run(p.add_run(page if page else ''), bold=False)


def add_manual_table_of_contents(doc):
    """Full static table of contents — always visible without Word field update."""
    entries = [
        (1, 'Abstract'),
        (1, '1. Introduction'),
        (2, '1.1 Background'),
        (2, '1.2 COOP Training Context'),
        (2, '1.3 Project Assignment'),
        (2, '1.4 Report Objectives'),
        (2, '1.5 Report Organization'),
        (1, '2. Training Environment and Organizational Context'),
        (2, '2.1 Employer Overview'),
        (2, '2.2 Role and Responsibilities'),
        (2, '2.3 Integration with Tawakkalna Services'),
        (1, '3. Project Scope and Requirements Analysis'),
        (2, '3.1 Users and Project Scope'),
        (2, '3.2 System Users'),
        (2, '3.3 Functional Requirements'),
        (2, '3.4 Non-Functional Requirements'),
        (1, '4. System Analysis and Design'),
        (2, '4.1 Use Case Diagram'),
        (2, '4.2 Sequence Diagram — Submit Review'),
        (2, '4.3 Entity Relationship Diagram'),
        (2, '4.4 System Architecture'),
        (1, '5. Technology Stack and Development Environment'),
        (1, '6. System Implementation'),
        (2, '6.1 Backend — Django REST Framework'),
        (2, '6.2 Frontend — React Visitor Portal'),
        (2, '6.3 Data Collection and Content Development'),
        (1, '7. Database Design'),
        (1, '8. Testing and Evaluation'),
        (2, '8.1 Testing Approach'),
        (2, '8.2 System Output'),
        (1, '9. Web System Interface'),
        (1, '10. Critical Analysis and Professional Reflection'),
        (2, '10.1 Most Challenging Tasks'),
        (2, '10.2 Tasks Performed Most Effectively'),
        (2, '10.3 Alternative Solutions Considered'),
        (2, '10.4 Professional Learning Outcomes'),
        (1, '11. Conclusions and Recommendations'),
        (2, '11.1 Conclusions'),
        (2, '11.2 Recommendations'),
        (1, 'References'),
        (1, 'Appendix A: REST API Endpoint Reference'),
        (1, 'Appendix B: Automated Test Command Output'),
    ]
    for level, title in entries:
        add_toc_line(doc, title, level=level)
    add_para(doc,
             'Note: Add page numbers in Word by placing the cursor after each dot line '
             'or use References → Table of Contents → Custom Table of Contents.',
             italic=True)


def add_toc_field(doc):
    """Insert Word TOC field — right-click → Update Field in Word to refresh page numbers."""
    add_para(doc,
             'Table of Contents', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
             )
    p = doc.add_paragraph()
    style_paragraph(p)
    run = p.add_run()
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    run._r.append(fld_begin)

    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = r'TOC \o "1-3" \h \z \u'
    run._r.append(instr)

    fld_sep = OxmlElement('w:fldChar')
    fld_sep.set(qn('w:fldCharType'), 'separate')
    run._r.append(fld_sep)

    run2 = p.add_run('[Open in Microsoft Word → right-click here → Update Field]')
    style_run(run2, italic=True)

    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    run._r.append(fld_end)
    doc.add_paragraph()


def add_page_numbers(doc):
    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.0)
        run = p.add_run()
        fld_begin = OxmlElement('w:fldChar')
        fld_begin.set(qn('w:fldCharType'), 'begin')
        run._r.append(fld_begin)
        instr = OxmlElement('w:instrText')
        instr.text = 'PAGE'
        run._r.append(instr)
        fld_sep = OxmlElement('w:fldChar')
        fld_sep.set(qn('w:fldCharType'), 'separate')
        run._r.append(fld_sep)
        run._r.append(OxmlElement('w:t'))
        fld_end = OxmlElement('w:fldChar')
        fld_end.set(qn('w:fldCharType'), 'end')
        run._r.append(fld_end)


def centered_block(doc, lines, size=FONT_SIZE, bold_first=False):
    for i, line in enumerate(lines):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.5)
        run = p.add_run(line)
        style_run(run, bold=(bold_first and i == 0), size=size)


# ── Preliminary pages ────────────────────────────────────────────────────────

def build_front_cover(doc):
    doc.add_paragraph()
    doc.add_paragraph()
    centered_block(doc, ['COOP Training Final Report'], size=Pt(22), bold_first=True)
    doc.add_paragraph()
    centered_block(doc, [REPORT_TITLE], size=Pt(16))
    doc.add_paragraph()
    doc.add_paragraph()
    centered_block(doc, [
        STUDENT_NAME,
        DEPARTMENT,
        f'Student ID: {STUDENT_ID}',
        f'Faculty Advisor: {FACULTY_ADVISOR}',
    ])
    add_page_break(doc)


def build_title_page(doc):
    centered_block(doc, [
        'University of Hail',
        'Kingdom of Saudi Arabia',
        '',
        DEPARTMENT,
        'COOP Education Program',
        '',
        REPORT_TITLE,
        '',
        f'Employer: {EMPLOYER}',
        f'Work Unit: {WORK_UNIT}',
        '',
        f'Prepared by: {STUDENT_NAME}',
        f'Student ID: {STUDENT_ID}',
        f'Faculty Advisor: {FACULTY_ADVISOR}',
        '',
        REPORT_DATE,
    ])
    add_page_break(doc)


def build_preliminary_lists(doc):
    add_heading(doc, 'Table of Contents', level=1)
    add_manual_table_of_contents(doc)
    add_page_break(doc)

    add_heading(doc, 'List of Tables', level=1)
    table_items = [
        ('Table 1', 'System Users and Roles'),
        ('Table 2', 'Functional Requirements Summary'),
        ('Table 3', 'Non-Functional Requirements Summary'),
        ('Table 4', 'Software Requirements'),
        ('Table 5', 'Main REST API Endpoints'),
        ('Table 6', 'Database Models Overview'),
        ('Table 7', 'Hardware Requirements'),
        ('Table 8', 'Testing Summary'),
        ('Table 9', 'Automated Unit Test Coverage'),
        ('Table 10', 'COOP Training Tasks and Outcomes'),
    ]
    for num, title in table_items:
        add_toc_line(doc, f'{num}: {title}', level=1)
    add_page_break(doc)

    add_heading(doc, 'List of Figures', level=1)
    figure_items = [
        ('Figure 1', 'Use Case Diagram — Jadah Hail Digital Platform'),
        ('Figure 2', 'Sequence Diagram — Submit Review Flow'),
        ('Figure 3', 'Entity Relationship Diagram (ERD)'),
        ('Figure 4', 'Three-Tier System Architecture'),
        ('Figure 5', 'Home Page — Visitor Portal'),
        ('Figure 6', 'Explore Page — Tourist Attractions'),
        ('Figure 7', 'Interactive Map — Hail Region'),
        ('Figure 8', 'Tourist Routes Page'),
        ('Figure 9', 'AI Tourism Assistant'),
        ('Figure 10', 'Place Details Page'),
        ('Figure 11', 'Admin Dashboard'),
    ]
    for num, title in figure_items:
        add_toc_line(doc, f'{num}: {title}', level=1)
    add_page_break(doc)

    add_heading(doc, 'List of Appendices', level=1)
    add_toc_line(doc, 'Appendix A: REST API Endpoint Reference', level=1)
    add_toc_line(doc, 'Appendix B: Automated Test Command Output', level=1)
    add_page_break(doc)


# ── Main body ────────────────────────────────────────────────────────────────

def build_abstract(doc):
    add_heading(doc, 'Abstract')
    add_para(doc,
        'This cooperative training (COOP) report documents the author\'s experiential learning '
        'at Hail Municipality, Technology Department, Digital Transformation Agency, where the '
        'primary assignment was to contribute to digital municipality services delivered through '
        'the Tawakkalna platform. The main project — Jadah Hail Digital (جادة حائل الرقمية) — '
        'is a smart tourism guide that enables visitors to explore Hail\'s landmarks, events, '
        'routes, and services through a bilingual web interface linked to municipality digital '
        'channels. The report describes requirements analysis, system design, full-stack '
        'implementation using Django REST Framework and React, data collection for tourist '
        'content, testing (19 automated unit tests, all passed), and professional evaluation '
        'of technical decisions. The experience strengthened skills in digital government '
        'service delivery, data organization, innovation under municipal standards, and '
        'cross-functional collaboration within a smart-city context.')
    add_page_break(doc)


def build_introduction(doc):
    add_heading(doc, '1. Introduction')

    add_para(doc, '1.1 Background', bold=True)
    add_para(doc,
        'Saudi Arabia\'s Vision 2030 emphasizes digital transformation of government services '
        'to improve citizen and visitor experience. Hail Municipality, through its Digital '
        'Transformation Agency, delivers municipal services via modern digital channels including '
        'Tawakkalna — the national unified platform for government and private services. '
        'Tourism promotion is a strategic priority for the Hail region, which hosts UNESCO-listed '
        'heritage sites, natural landmarks, and cultural festivals. A structured digital tourist '
        'guide accessible within municipality services on Tawakkalna addresses the need for '
        'accurate, centralized, and user-friendly tourism information.')

    add_para(doc, '1.2 COOP Training Context', bold=True)
    add_para(doc,
        f'During the cooperative training period at {EMPLOYER}, the author was assigned to '
        f'{WORK_UNIT}. The training followed a structured progression: an initial orientation '
        'and qualification phase focused on understanding municipal IT governance, digital '
        'service standards, and the Tawakkalna services ecosystem; followed by direct technical '
        'responsibilities including requirements gathering, data collection and documentation, '
        'system development, report preparation, and quality assurance. This report focuses on '
        'the author\'s individual contributions rather than general departmental procedures.')

    add_para(doc, '1.3 Project Assignment', bold=True)
    add_para(doc,
        'The assigned project was to develop Jadah Hail Digital — a smart digital tourist guide '
        'integrated as a municipality service within the Tawakkalna digital services framework. '
        'The platform (internally developed as a full-stack web application) provides visitors '
        'with searchable tourist places, interactive maps, event listings, suggested routes, '
        'user reviews, a rule-based tourism assistant, and an administrative interface for '
        'municipality staff. The service is designed to be accessible through Tawakkalna-linked '
        'municipality digital channels, extending Hail Municipality\'s smart-city tourism offering.')

    add_para(doc, '1.4 Report Objectives', bold=True)
    add_bullets(doc, [
        'Document the problem definition, requirements, and solution design for Jadah Hail Digital.',
        'Explain how the system was implemented and validated through testing.',
        'Present critical professional evaluation of technical choices and challenges encountered.',
        'Summarize learning outcomes and recommendations for future improvement.',
    ])

    add_para(doc, '1.5 Report Organization', bold=True)
    add_para(doc,
        'Section 2 describes the training environment and organizational context. '
        'Section 3 covers project scope and user requirements. '
        'Sections 4–8 address system design, technology stack, database, and implementation. '
        'Section 9 presents testing and results. Section 10 shows the web interface. '
        'Section 11 provides critical analysis and professional reflection. '
        'Section 12 concludes with recommendations. References and appendices follow.')


def build_organization(doc):
    add_heading(doc, '2. Training Environment and Organizational Context')

    add_para(doc, '2.1 Employer Overview', bold=True)
    add_para(doc,
        'Hail Municipality governs the Hail region in northwestern Saudi Arabia. Its Technology '
        'Department supports ICT infrastructure, application development, and digital service '
        'delivery. The Digital Transformation Agency within this department leads initiatives '
        'aligned with national digital government programs, including services published through '
        'Tawakkalna.')

    add_para(doc, '2.2 Role and Responsibilities', bold=True)
    add_table(doc,
        ['Phase', 'Author\'s Activities', 'Skills Developed'],
        [
            ['Orientation & Qualification',
             'Learning municipal IT policies, Tawakkalna service model, project scope',
             'Organizational awareness, digital governance'],
            ['Requirements & Data Collection',
             'Documenting tourist sites, events, routes; bilingual content preparation',
             'Data accuracy, field documentation, reporting'],
            ['Development',
             'Backend API, frontend portal, map integration, assistant module',
             'Full-stack development, API design, UX for public services'],
            ['Testing & Reporting',
             'Manual and automated testing, COOP and technical reports',
             'Quality assurance, technical writing, professional evaluation'],
        ],
        caption='Table 10: COOP Training Tasks and Outcomes')

    add_para(doc, '2.3 Integration with Tawakkalna Services', bold=True)
    add_para(doc,
        'Jadah Hail Digital was developed as a municipality digital tourism service intended for '
        'delivery through Tawakkalna\'s government services channel. Tawakkalna serves as the '
        'national access point where citizens and visitors reach Hail Municipality services; '
        'the tourist guide extends this ecosystem by providing structured tourism content, '
        'location data, and interactive exploration tools. The web application architecture '
        '(React frontend + Django REST API) supports deployment as a linked digital service '
        'within the municipality\'s Tawakkalna service portfolio, ensuring consistency with '
        'broader digital transformation objectives.')


def build_scope_and_requirements(doc):
    add_heading(doc, '3. Project Scope and Requirements Analysis')

    add_para(doc, '3.1 Users and Project Scope', bold=True)
    add_para(doc,
        'Jadah Hail Digital serves multiple stakeholder groups within the municipality\'s '
        'digital services framework.')
    add_bullets(doc, [
        'Visitor portal (Arabic/English, RTL/LTR) for exploring Hail tourism content.',
        'Interactive map (OpenStreetMap + Leaflet) with accurate landmark coordinates.',
        'Events, routes, reviews, favorites, and rule-based tourism assistant.',
        'Admin dashboard for municipality staff (statistics and content overview).',
        'REST API (Django REST Framework) supporting the portal and future integrations.',
    ])
    add_para(doc, 'Out of scope: native mobile apps, payment/booking, generative AI assistant.')

    add_para(doc, '3.2 System Users', bold=True)
    add_table(doc,
        ['User Type', 'Description', 'Permissions'],
        [
            ['Guest Visitor', 'Accesses service via Tawakkalna-linked portal',
             'Browse places, map, events, routes, assistant'],
            ['Registered Visitor', 'Authenticated user account',
             'Guest permissions + favorites + reviews'],
            ['Municipality Admin', 'Hail Municipality staff',
             'Admin dashboard + Django Admin'],
            ['IT Developer', 'System maintenance',
             'Server, database, deployment'],
        ],
        caption='Table 1: System Users and Roles')

    add_para(doc, '3.3 Functional Requirements', bold=True)
    add_table(doc,
        ['ID', 'Requirement'],
        [
            ['FR-01', 'Display tourist sites with category classification'],
            ['FR-02', 'Search and filter places by category'],
            ['FR-03', 'Detail page per landmark (description, hours, map, reviews)'],
            ['FR-04', 'Interactive Hail region map with markers'],
            ['FR-05', 'Display tourism events and suggested routes'],
            ['FR-06', 'User login/logout and session management'],
            ['FR-07', 'Favorites and review submission (1–5 rating)'],
            ['FR-08', 'Rule-based tourism assistant'],
            ['FR-09', 'Admin dashboard for content and analytics overview'],
            ['FR-10', 'Bilingual Arabic/English interface with RTL support'],
        ],
        caption='Table 2: Functional Requirements Summary')

    add_para(doc, '3.4 Non-Functional Requirements', bold=True)
    add_table(doc,
        ['ID', 'Requirement'],
        [
            ['NFR-01', 'Responsive, accessible user interface'],
            ['NFR-02', 'API response under 2 seconds (local environment)'],
            ['NFR-03', 'Session security and CSRF protection'],
            ['NFR-04', 'Client-server separation for scalability'],
            ['NFR-05', 'Extensibility for PostgreSQL and cloud deployment'],
        ],
        caption='Table 3: Non-Functional Requirements Summary')


def build_design(doc):
    add_heading(doc, '4. System Analysis and Design')

    add_para(doc, '4.1 Use Case Diagram', bold=True)
    add_para(doc,
        'Figure 1 models interactions between Visitor, Registered User, and Admin actors '
        'with the Jadah Hail Digital system.')
    add_image(doc, DIAGRAMS / 'use_case_diagram.png',
              caption='Figure 1: Use Case Diagram — Jadah Hail Digital Platform')

    add_para(doc, '4.2 Sequence Diagram — Submit Review', bold=True)
    add_para(doc,
        'Figure 2 illustrates the authenticated review submission flow between the React '
        'frontend, Django REST API, and SQLite database.')
    add_image(doc, DIAGRAMS / 'sequence_diagram.png',
              caption='Figure 2: Sequence Diagram — Submit Review Flow')

    add_para(doc, '4.3 Entity Relationship Diagram', bold=True)
    add_image(doc, DIAGRAMS / 'erd_diagram.png',
              caption='Figure 3: Entity Relationship Diagram (ERD)')

    add_para(doc, '4.4 System Architecture', bold=True)
    add_para(doc,
        'Figure 4 shows the three-tier architecture: Presentation (React), Application '
        '(Django REST Framework), and Data (SQLite + media files).')
    add_image(doc, DIAGRAMS / 'architecture_diagram.png',
              caption='Figure 4: Three-Tier System Architecture')


def build_technology(doc):
    add_heading(doc, '5. Technology Stack and Development Environment')

    add_table(doc,
        ['Software', 'Version', 'Purpose'],
        [
            ['Python', '3.13+', 'Backend language'],
            ['Django', '6.x', 'Web framework'],
            ['Django REST Framework', '3.15+', 'REST API'],
            ['React + TypeScript', '18.3 / 5.x', 'Visitor portal UI'],
            ['Vite + Tailwind CSS', '6.3 / 4.1', 'Build tooling and styling'],
            ['Leaflet + OpenStreetMap', '1.9', 'Interactive map'],
            ['SQLite', '3', 'Development database'],
            ['Git / GitHub', '2.x', 'Version control'],
        ],
        caption='Table 4: Software Requirements')

    add_para(doc, 'Development environment: Windows 10/11, VS Code/Cursor IDE. '
             'Backend: python manage.py runserver (port 8000). '
             'Frontend: npm run dev (port 5173). '
             'Repository: https://github.com/moneerafahaid-collab/WATN-2-.git')

    add_table(doc,
        ['Component', 'Minimum', 'Recommended'],
        [
            ['Processor', 'Intel i3 / Ryzen 3', 'Intel i5 / Ryzen 5'],
            ['RAM', '8 GB', '16 GB'],
            ['Storage', '10 GB free', '20 GB SSD'],
            ['Network', 'Internet required', 'Broadband'],
        ],
        caption='Table 7: Hardware Requirements')


def build_backend_frontend(doc):
    add_heading(doc, '6. System Implementation')

    add_para(doc, '6.1 Backend — Django REST Framework', bold=True)
    add_para(doc, 'The backend is organized into Django apps:')
    add_bullets(doc, [
        'tourism: Tourist places, reviews, routes.',
        'events: Tourism events.',
        'accounts: Profiles, favorites, session authentication.',
        'ai_assistant: Rule-based tourism assistant replies.',
    ])
    add_table(doc,
        ['Endpoint', 'Method', 'Description'],
        [
            ['/places/', 'GET', 'List tourist places'],
            ['/reviews/', 'GET/POST', 'View/add reviews'],
            ['/routes/', 'GET', 'Tourist routes'],
            ['/events/', 'GET', 'Events'],
            ['/favorites/', 'GET/POST/DELETE', 'Favorites'],
            ['/auth/login/', 'POST', 'Login'],
            ['/auth/logout/', 'POST', 'Logout'],
            ['/assistant/ask/', 'POST', 'Tourism assistant query'],
        ],
        caption='Table 5: Main REST API Endpoints')

    add_para(doc, '6.2 Frontend — React Visitor Portal', bold=True)
    add_bullets(doc, [
        'TourismContext for centralized API data management.',
        'HailMap component (Leaflet) for geographic visualization.',
        'locale.ts for Arabic RTL and English LTR localization.',
        'Dual-mode interface: Visitor Portal and Admin Dashboard.',
        'Vite dev proxy routes /api and /media to Django backend.',
    ])

    add_para(doc, '6.3 Data Collection and Content Development', bold=True)
    add_para(doc,
        'A significant COOP responsibility was collecting and verifying tourism data for Hail '
        'landmarks, including bilingual names and descriptions, geographic coordinates, visiting '
        'hours, and authentic images sourced from Wikimedia Commons. Management commands '
        '(load_demo_data, sync_arabic_content, import_hail_images) automate content loading. '
        'This field work ensured the digital guide reflects accurate municipal tourism information '
        'suitable for public service delivery via Tawakkalna-linked channels.')


def build_database(doc):
    add_heading(doc, '7. Database Design')

    add_table(doc,
        ['Model', 'Key Fields', 'Relationships'],
        [
            ['TouristPlace', 'name_ar/en, category, lat/lng, image',
             '← Review; M2M Route, Favorites'],
            ['Review', 'user, place, rating, comment', 'FK → User, TouristPlace'],
            ['TouristRoute', 'name, duration, difficulty', 'M2M → TouristPlace'],
            ['Event', 'title_ar/en, dates, location', 'Standalone'],
            ['UserProfile', 'language, interests, favorites', 'OneToOne → User'],
            ['AssistantReply', 'prompt/response ar/en, keywords', 'Standalone'],
        ],
        caption='Table 6: Database Models Overview')


def build_testing(doc):
    add_heading(doc, '8. Testing and Evaluation')

    add_para(doc, '8.1 Testing Approach', bold=True)
    add_para(doc,
        'Testing combined manual functional verification with automated Django unit tests '
        'to ensure API reliability before municipal service deployment.')

    add_table(doc,
        ['Test Type', 'Method', 'Result'],
        [
            ['Manual functional testing', 'All portal pages and admin UI', 'Passed'],
            ['API integration testing', 'Browser DevTools / direct API calls', 'Passed'],
            ['Authentication flow', 'Login, logout, session persistence', 'Passed'],
            ['Map rendering', 'Leaflet markers on OpenStreetMap tiles', 'Passed'],
            ['Localization', 'Arabic RTL / English LTR switching', 'Passed'],
            ['Automated unit tests', 'python manage.py test (19 tests)', 'Passed — 19/19'],
        ],
        caption='Table 8: Testing Summary')

    add_table(doc,
        ['Module', 'Tests', 'Coverage'],
        [
            ['tourism.tests', '10', 'Places, reviews, routes API'],
            ['accounts.tests', '7', 'Auth, session, favorites'],
            ['events.tests', '2', 'Events list, bilingual fields'],
        ],
        caption='Table 9: Automated Unit Test Coverage')

    add_para(doc, '8.2 System Output', bold=True)
    add_para(doc,
        'The deployed system presents 6 tourist landmarks, 3 events, 3 routes, a tourism '
        'assistant with stored replies, and a bilingual admin dashboard — fulfilling the '
        'functional scope defined for Jadah Hail Digital as a Tawakkalna-linked municipality service.')


def build_interface(doc):
    add_heading(doc, '9. Web System Interface')

    add_para(doc,
        'Figures 5–11 present screenshots of Jadah Hail Digital captured from the live '
        'development environment during the COOP training period.')

    for filename, caption in [
        ('01_home.png', 'Figure 5: Home Page — Visitor Portal'),
        ('02_explore.png', 'Figure 6: Explore Page — Tourist Attractions'),
        ('03_map.png', 'Figure 7: Interactive Map — Hail Region'),
        ('04_routes.png', 'Figure 8: Tourist Routes Page'),
        ('05_assistant.png', 'Figure 9: AI Tourism Assistant'),
        ('06_place_details.png', 'Figure 10: Place Details Page'),
        ('07_admin_dashboard.png', 'Figure 11: Admin Dashboard'),
    ]:
        add_image(doc, SCREENSHOTS / filename, caption=caption)


def build_critical_analysis(doc):
    add_heading(doc, '10. Critical Analysis and Professional Reflection')

    add_para(doc,
        'Per COOP report guidelines, this section evaluates the author\'s experiential learning '
        'and professional judgment regarding the Jadah Hail Digital project.')

    add_para(doc, '10.1 Most Challenging Tasks', bold=True)
    add_bullets(doc, [
        'Bilingual content management (Arabic RTL + English LTR) required careful UI design '
        'and duplicate database fields — more complex than monolingual applications because '
        'layout direction, font selection, and content synchronization had to remain consistent.',
        'Geographic data accuracy for the interactive map demanded verified coordinates for '
        'each landmark; incorrect data would undermine public trust in a municipality service.',
        'Aligning the standalone web application with Tawakkalna service delivery expectations '
        'required understanding both technical API design and municipal digital governance.',
    ])

    add_para(doc, '10.2 Tasks Performed Most Effectively', bold=True)
    add_bullets(doc, [
        'REST API design with Django REST Framework — modular ViewSets and serializers '
        'enabled clean separation between frontend and backend, simplifying future Tawakkalna '
        'integration and mobile extension.',
        'Automated data loading via management commands — reduced manual errors when populating '
        'tourist content and supported reproducible deployment.',
        'Structured unit testing (19 tests) — provided confidence in authentication, reviews, '
        'and favorites logic critical to registered visitor features.',
    ])

    add_para(doc, '10.3 Alternative Solutions Considered', bold=True)
    add_para(doc,
        'A monolithic Django template approach could have reduced stack complexity but would '
        'limit UI richness and Tawakkalna service embedding flexibility. A no-code CMS was '
        'considered insufficient for interactive maps, custom assistant logic, and API-driven '
        'architecture required by the Digital Transformation Agency. The chosen React + Django '
        'split balances modern UX with robust backend governance — appropriate for municipal '
        'digital services despite higher initial setup cost.')

    add_para(doc, '10.4 Professional Learning Outcomes', bold=True)
    add_bullets(doc, [
        'Organization: Understanding municipal IT workflows and report-driven accountability.',
        'Innovation: Applying modern web technologies within government service constraints.',
        'Mastery: Full-stack development from requirements through testing and documentation.',
        'Digital governance: Aligning project deliverables with Tawakkalna service ecosystem goals.',
    ])


def build_conclusion(doc):
    add_heading(doc, '11. Conclusions and Recommendations')

    add_para(doc, '11.1 Conclusions', bold=True)
    add_para(doc,
        'The COOP training at Hail Municipality\'s Digital Transformation Agency provided '
        'practical experience in delivering Jadah Hail Digital — a smart tourism guide integrated '
        'within Tawakkalna municipality services. The author contributed to requirements analysis, '
        'data collection, full-stack development, testing, and reporting. The resulting platform '
        'demonstrates a scalable client-server architecture with bilingual support, interactive '
        'mapping, and validated API endpoints. Nineteen automated unit tests passed, confirming '
        'core functionality. The project aligns with Vision 2030 digital government objectives '
        'and enhances Hail\'s tourism promotion capabilities.')

    add_para(doc, '11.2 Recommendations', bold=True)
    add_bullets(doc, [
        'Complete deep integration with Tawakkalna service APIs for seamless in-app access.',
        'Connect admin dashboard CRUD operations directly to the REST API.',
        'Migrate from SQLite to PostgreSQL for production municipal deployment.',
        'Expand automated testing to frontend components (React Testing Library).',
        'Develop a companion mobile experience consuming the same REST API.',
        'Evaluate LLM-based assistant upgrade while maintaining municipal content control.',
    ])


def build_references(doc):
    add_heading(doc, 'References')
    refs = [
        'Django Software Foundation. (2025). Django Documentation. https://docs.djangoproject.com/',
        'Django REST Framework. (2025). API Guide. https://www.django-rest-framework.org/',
        'Meta Open Source. (2025). React Documentation. https://react.dev/',
        'Vite Team. (2025). Vite Guide. https://vite.dev/',
        'Leaflet. (2025). Interactive maps library. https://leafletjs.com/',
        'OpenStreetMap contributors. (2025). https://www.openstreetmap.org/',
        'UNESCO. Rock Art in the Hail Region of Saudi Arabia. https://whc.unesco.org/',
        'Kingdom of Saudi Arabia. (2023). Vision 2030 — Digital Government. https://www.vision2030.gov.sa/',
        'University of Hail, CCSE. COOP Report Writing Guidelines.',
        'Wikimedia Commons. Hail region landmark images. https://commons.wikimedia.org/',
    ]
    for i, ref in enumerate(refs, 1):
        add_para(doc, f'[{i}] {ref}')


def build_appendices(doc):
    add_page_break(doc)
    add_heading(doc, 'Appendix A: REST API Endpoint Reference')
    add_table(doc,
        ['Endpoint', 'Method', 'Auth', 'Description'],
        [
            ['/api/v1/places/', 'GET', 'None', 'List/filter tourist places'],
            ['/api/v1/places/{id}/', 'GET', 'None', 'Place detail'],
            ['/api/v1/reviews/', 'GET/POST', 'POST: Yes', 'Reviews list/create'],
            ['/api/v1/routes/', 'GET', 'None', 'Tourist routes'],
            ['/api/v1/events/', 'GET', 'None', 'Events list'],
            ['/api/v1/favorites/', 'GET/POST/DELETE', 'Yes', 'User favorites'],
            ['/api/v1/auth/login/', 'POST', 'None', 'Session login'],
            ['/api/v1/auth/logout/', 'POST', 'Yes', 'Session logout'],
            ['/api/v1/auth/me/', 'GET', 'None', 'Current user status'],
            ['/api/v1/assistant/ask/', 'POST', 'None', 'Assistant query'],
        ])

    add_page_break(doc)
    add_heading(doc, 'Appendix B: Automated Test Command Output')
    add_para(doc,
        'Command executed: python manage.py test tourism accounts events\n\n'
        'Result: Ran 19 tests — OK.\n\n'
        'Modules tested: tourism.tests (10), accounts.tests (7), events.tests (2).')


def build_document():
    doc = Document()
    for section in doc.sections:
        section.page_height = Cm(29.7)
        section.page_width = Cm(21)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    build_front_cover(doc)
    build_title_page(doc)
    build_preliminary_lists(doc)
    build_abstract(doc)
    build_introduction(doc)
    build_organization(doc)
    build_scope_and_requirements(doc)
    build_design(doc)
    build_technology(doc)
    build_backend_frontend(doc)
    build_database(doc)
    build_testing(doc)
    build_interface(doc)
    build_critical_analysis(doc)
    build_conclusion(doc)
    build_references(doc)
    build_appendices(doc)

    add_page_numbers(doc)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == '__main__':
    import subprocess
    import sys

    subprocess.run([sys.executable, str(ROOT / 'scripts' / 'generate_diagrams.py')], check=True)
    if not any(SCREENSHOTS.glob('*.png')):
        print('Note: Run capture_screenshots.py with servers running for interface figures.')
    path = build_document()
    print(f'Created: {path}')
    print('Table of Contents is pre-filled. Add page numbers in Word if required by your advisor.')
